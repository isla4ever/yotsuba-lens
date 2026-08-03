from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "store-assets"
SOURCE = ROOT / "docs" / "assets"
ASSETS.mkdir(parents=True, exist_ok=True)
PUBLIC_ICON = ROOT / "public" / "icon"
PUBLIC_ICON.mkdir(parents=True, exist_ok=True)

BG = "#f5f6f0"
INK = "#111812"
MUTED = "#4f5a51"
LIME = "#84ad13"
LIME_DARK = "#5f7f08"
DARK = "#101712"
LIGHT = "#f5f3e9"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def ai_asset(name):
    path = ASSETS / "source" / name
    if not path.exists():
        raise FileNotFoundError(f"Missing required gpt-image-2 source asset: {path}")
    return Image.open(path).convert("RGB")


def cover_crop(image, width, height, horizontal="center", vertical="center"):
    target_ratio = width / height
    source_ratio = image.width / image.height
    if source_ratio > target_ratio:
        crop_width = round(image.height * target_ratio)
        left = 0 if horizontal == "left" else (image.width - crop_width) // 2
        box = (left, 0, left + crop_width, image.height)
    else:
        crop_height = round(image.width / target_ratio)
        top = 0 if vertical == "top" else (image.height - crop_height) // 2
        box = (0, top, image.width, top + crop_height)
    return image.crop(box).resize((width, height), Image.Resampling.LANCZOS)


def make_icon():
    source = ai_asset("yotsuba-lens-logo-final-gpt-image-2.png")
    square_size = min(source.width, source.height)
    left = (source.width - square_size) // 2
    top = (source.height - square_size) // 2
    source = source.crop((left, top, left + square_size, top + square_size))
    for size in [16, 32, 48, 128]:
        source.resize((size, size), Image.Resampling.LANCZOS).save(PUBLIC_ICON / f"{size}.png", format="PNG", optimize=True)
    Image.open(PUBLIC_ICON / "128.png").save(ASSETS / "icon-128.png", format="PNG", optimize=True)


def make_small_promo():
    source = ai_asset("yotsuba-lens-promo-small-final-gpt-image-2.png")
    image = cover_crop(source, 440, 280, horizontal="left")
    image.save(ASSETS / "promo-small-440x280.png", format="PNG", optimize=True)


def make_top_promo():
    source = ai_asset("yotsuba-lens-promo-top-final-gpt-image-2.png")
    image = cover_crop(source, 1400, 560, vertical="top")
    image.save(ASSETS / "promo-top-1400x560.png", format="PNG", optimize=True)


def normalize_screenshots():
    def browser_frame(draw, x, y, width, height, title):
        draw.rounded_rectangle((x, y, x + width, y + height), radius=14, fill="#ffffff", outline="#d8ded4", width=2)
        draw.rectangle((x, y, x + width, y + 40), fill="#edf0e9")
        for dot_x, color in [(x + 18, "#ef7567"), (x + 36, "#e7bb55"), (x + 54, "#76b85d")]:
            draw.ellipse((dot_x - 5, y + 15, dot_x + 5, y + 25), fill=color)
        draw.rounded_rectangle((x + 84, y + 10, x + width - 18, y + 30), radius=8, fill="#ffffff", outline="#d6ddd2")
        draw.text((x + 98, y + 13), title, fill=MUTED, font=font(12))

        draw.text((x + 28, y + 72), "Atlas Studio", fill=INK, font=font(28, True))
        draw.text((x + 28, y + 112), "Product design workspace", fill=MUTED, font=font(15))
        draw.rounded_rectangle((x + 28, y + 150, x + width - 28, y + 246), radius=10, fill="#f0f6d8")
        draw.text((x + 48, y + 176), "把复杂页面变成清晰证据", fill=INK, font=font(20, True))
        for index, label in enumerate(["结构", "状态", "动效"]):
            card_x = x + 28 + index * ((width - 70) // 3)
            card_w = (width - 92) // 3
            draw.rounded_rectangle((card_x, y + 274, card_x + card_w, y + 408), radius=8, fill="#fafbf8", outline="#e0e5dd")
            draw.text((card_x + 16, y + 296), label, fill=LIME_DARK, font=font(14, True))
            draw.rectangle((card_x + 16, y + 332, card_x + card_w - 16, y + 340), fill="#dfe7d7")
            draw.rectangle((card_x + 16, y + 354, card_x + card_w - 34, y + 362), fill="#e8ede5")
        draw.rounded_rectangle((x + 28, y + 444, x + width - 28, y + height - 28), radius=8, fill="#f7f9f5")
        draw.text((x + 48, y + 470), "可复用的视觉规律", fill=INK, font=font(18, True))
        for index in range(4):
            draw.rectangle((x + 48, y + 510 + index * 34, x + width - 72 - index * 34, y + 519 + index * 34), fill="#dfe6db")

    def compose(name, panel_candidates, caption, accent):
        canvas = Image.new("RGB", (1280, 800), BG)
        draw = ImageDraw.Draw(canvas)
        draw.rectangle((0, 0, 1280, 12), fill=accent)
        draw.text((46, 28), "Yotsuba Lens", fill=INK, font=font(24, True))
        draw.text((46, 62), caption, fill=MUTED, font=font(15))
        browser_frame(draw, 42, 112, 734, 640, "atlas.example/design-system")
        panel = None
        for candidate in panel_candidates:
            if candidate.exists():
                panel = Image.open(candidate).convert("RGB")
                break
        if panel is None:
            panel = Image.new("RGB", (360, 800), "#f5f6f0")
        panel = panel.crop((0, 0, panel.width, min(panel.height, 800)))
        panel.thumbnail((410, 744), Image.Resampling.LANCZOS)
        panel_x = 828 + (410 - panel.width) // 2
        panel_y = 34 + (744 - panel.height) // 2
        draw.rounded_rectangle((810, 22, 1250, 778), radius=16, fill="#dfe4db", outline="#c8d0c4", width=2)
        canvas.paste(panel, (panel_x, panel_y))
        canvas.save(ASSETS / name, format="PNG", optimize=True)

    compose(
        "screenshot-evidence-workspace-1280x800.png",
        [ROOT / "output" / "playwright" / "extension-ui" / "sidepanel-settings-zh-light.png", ROOT / "output" / "playwright" / "extension-ui" / "sidepanel-zh-light.png"],
        "侧边栏设置：模式、输出与 AI 配置集中管理",
        LIME,
    )
    compose(
        "screenshot-smart-capture-1280x800.png",
        [ROOT / "output" / "playwright" / "extension-ui" / "popup-zh-light.png", ROOT / "output" / "playwright" / "extension-ui" / "sidepanel-first-use-ai-guide.png"],
        "智能捕获：一次操作开始，页面保持响应",
        "#c8ff56",
    )


def add_paragraph(doc, text="", bold=False, color=None, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    set_docx_font(r)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def set_docx_font(run):
    # Keep the East Asian family explicit so Word and macOS Preview select a
    # Chinese-capable face instead of inheriting the Latin body font.
    family = "Hiragino Sans GB"
    run.font.name = family
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), family)
    r_fonts.set(qn("w:hAnsi"), family)
    r_fonts.set(qn("w:eastAsia"), family)
    r_fonts.set(qn("w:cs"), family)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def add_field(doc, label, value, note=""):
    add_paragraph(doc, label, bold=True, color="5F7F08", size=11, after=2)
    add_paragraph(doc, value, size=11, after=2)
    if note:
        add_paragraph(doc, "后台提示：" + note, color="555555", size=9, after=8)


def make_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Yotsuba 网页风格提取器 Chrome 商店中文上架填写包")
    set_docx_font(run)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(17, 24, 18)
    add_paragraph(doc, "版本：0.3.1　｜　仅上传标准版　｜　发布范围：公开　｜　审核通过后暂存", color="5F7F08", size=10, after=14)
    add_paragraph(doc, "使用说明", bold=True, color="111812", size=15, after=6)
    add_paragraph(doc, "按 Chrome Web Store 开发者后台从上到下填写。文字字段可直接复制；图片文件位于项目 docs/store-assets/ 目录。首次提交建议关闭自动发布，审核通过后再手动公开。", size=10, after=12)

    add_paragraph(doc, "一、商品详情", bold=True, color="111812", size=15, after=8)
    add_field(doc, "软件包中的标题", "Yotsuba 网页风格提取器", "保持与新标准版扩展包 manifest 中的名称一致。")
    add_field(doc, "软件包中的摘要", "提取当前网页的配色、布局、组件、交互和动效，生成可直接使用的 Prompt，帮助你参考或还原网页风格。", "上传新的标准版 ZIP 后由软件包自动带入。")
    add_field(doc, "说明", "看到喜欢的网页，想参考它的配色、布局或交互，却不知道该怎样准确描述？Yotsuba 网页风格提取器可以把当前页面整理成清晰的风格参考和可直接使用的 Prompt。\n\n打开网页后点击“智能捕获”，即可提取页面的配色、字体、间距、布局、组件、交互状态和动效。捕获结果会显示在 Chrome 侧边栏中，你可以查看已整理的内容、缺少的状态和下一步建议。\n\n主要功能：\n• 一键提取网页配色、布局、组件、交互和动效\n• 生成适合常见编码工具使用的网页实现 Prompt\n• “设计参照”模式用于借鉴网页风格并设计原创页面\n• “经授权重建”模式用于整理明确页面、尺寸和状态的还原资料\n• 自动提示需要补充的悬停、展开、滚动或响应式状态\n• 导出 Prompt、风格参考和重建资料包\n\n基础提取不需要账号，也不需要配置模型服务。采集、历史和导出默认在本机完成；只有你主动配置服务并要求生成内容时，才会发送精简后的页面特征。\n\n扩展不会在后台自动扫描网页，不会自动点击、输入、提交表单或跳转页面。请只在你有权查看、参考或还原的网页上使用。", "完整粘贴到说明文本框。")
    add_field(doc, "类别", "开发者工具", "选择 Developer Tools / 开发者工具。")
    add_field(doc, "语言", "中文（简体）", "选择 Chinese (Simplified) / 中文（简体）。")

    add_paragraph(doc, "二、图片资源", bold=True, color="111812", size=15, after=8)
    add_field(doc, "商店图标", "docs/store-assets/icon-128.png", "128×128，24 位 PNG，无透明层。")
    add_field(doc, "屏幕截图 1", "docs/store-assets/screenshot-overview-dark-1280x800.png", "1280×800，夜间模式概览，24 位 PNG，无透明层。")
    add_field(doc, "屏幕截图 2", "docs/store-assets/screenshot-coverage-dark-1280x800.png", "1280×800，夜间模式覆盖，24 位 PNG，无透明层。")
    add_field(doc, "屏幕截图 3", "docs/store-assets/screenshot-history-dark-1280x800.png", "1280×800，夜间模式历史，24 位 PNG，无透明层。")
    add_field(doc, "屏幕截图 4", "docs/store-assets/screenshot-settings-dark-1280x800.png", "1280×800，夜间模式设置，24 位 PNG，无透明层。")
    add_field(doc, "小型宣传图块", "docs/store-assets/promo-small-440x280.png", "440×280，24 位 PNG，无透明层。")
    add_field(doc, "顶部宣传图块", "docs/store-assets/promo-top-1400x560.png", "1400×560，24 位 PNG，无透明层。")
    add_field(doc, "宣传视频", "留空", "项目当前没有公开视频，不要填写无关链接。")

    add_paragraph(doc, "三、其他字段", bold=True, color="111812", size=15, after=8)
    add_field(doc, "官方网站", "https://github.com/isla4ever/yotsuba-lens", "选择 GitHub 项目主页。")
    add_field(doc, "主页网址", "https://github.com/isla4ever/yotsuba-lens", "如果后台把官方网站与主页分开，两个字段都填项目主页。")
    add_field(doc, "支持信息页面网址", "https://github.com/isla4ever/yotsuba-lens/issues", "用于接收问题反馈和使用支持。")
    add_field(doc, "成人内容", "否 / 关闭", "Yotsuba Lens 不包含成人内容。")
    add_field(doc, "商品支持", "公开范围 / 开启", "允许用户从商店详情页进入支持入口；支持地址使用 GitHub Issues。")

    add_paragraph(doc, "截图对应的填写顺序", bold=True, color="111812", size=15, after=8)
    for item in [
        "软件包中的标题：保留 Yotsuba 网页风格提取器。",
        "软件包中的摘要：保留上传 ZIP 带入的摘要，或粘贴上方中文摘要。",
        "说明：粘贴上方完整中文说明。",
        "类别：选择开发者工具。语言：选择中文（简体）。",
        "商店图标：上传 icon-128.png。宣传视频：留空。",
        "屏幕截图：按概览、覆盖、历史、设置的顺序上传四张 1280×800 截图。",
        "小型宣传图块：上传 promo-small-440x280.png；顶部宣传图块：上传 promo-top-1400x560.png。",
        "官方网站、主页网址、支持信息页面网址：按上方地址填写；成人内容关闭；商品支持选择公开范围并开启。",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_docx_font(r)
        r.font.size = Pt(10)

    add_paragraph(doc, "四、隐私权与审核", bold=True, color="111812", size=15, after=8)
    add_field(doc, "隐私权政策网址", "https://github.com/isla4ever/yotsuba-lens/blob/main/docs/privacy.md", "必须使用可公开访问的 HTTPS 地址。")
    add_field(doc, "单一用途", "Yotsuba 网页风格提取器的单一用途是：在用户主动操作时提取当前网页的配色、布局、组件、交互、动效和截图，并整理成网页风格参考、Prompt 或经授权的还原资料。", "按后台提示完整填写，不要只写产品名称。")
    add_field(doc, "activeTab 权限理由", "仅在用户点击扩展按钮、智能捕获或选取组件时临时访问当前活动标签页，以读取当前页面并完成用户明确发起的采集。扩展不会在后台自动访问其他标签页。")
    add_field(doc, "scripting 权限理由", "当用户主动开始采集时，按需向当前标签页注入随扩展包发布的本地采集脚本，用于读取页面结构、计算样式和交互状态。不会注入或执行任何远程代码。")
    add_field(doc, "storage 权限理由", "在 Chrome 本地存储中保存界面语言、主题、采集设置、用户配置的 AI 服务资料、捕获历史、证据元数据和重建项目草稿，以便用户恢复工作区并导出证据。数据默认仅保存在本机。")
    add_field(doc, "tabs 权限理由", "用于识别用户当前活动标签页、读取其 URL 和标题、在标签页切换时刷新侧边栏状态，并在采集期间与当前页面通信和获取当前可见区域截图。不会用于跟踪与采集无关的浏览活动。")
    add_field(doc, "sidePanel 权限理由", "用于将捕获模式、采集进度、证据覆盖、历史记录和设置展示在 Chrome 原生侧边栏中，使用户在不离开当前网页的情况下操作和检查采集结果。")
    add_field(doc, "主机权限理由", "用户可能需要在其有权分析的任意 HTTP 或 HTTPS 网页上采集设计证据，因此扩展需要访问用户当前选择的网站。访问只在用户主动发起捕获、选取组件或补采状态时发生；不会自动扫描所有网站，chrome:// 等受限页面会被拒绝。")
    add_field(doc, "远程代码", "选择：不，我并未使用远程代码", "所有 JavaScript 依赖和采集逻辑均随扩展包发布。可选 AI 接口只交换数据和文本结果，不下载或执行远程代码。")
    add_field(doc, "远程代码理由（仅在输入框仍显示时填写）", "本扩展不使用远程代码。所有 JavaScript 依赖和采集逻辑均已打包在提交的软件包中，不通过 eval、new Function、远程 script 或远程模块执行代码。可选 AI 接口只交换结构化数据和文本结果，不下载或执行服务端返回的代码。")
    add_field(doc, "数据类型勾选", "勾选：身份验证信息、网络记录、用户活动、网站内容。不要勾选：个人身份信息、健康信息、财务和付款信息、个人通讯、位置。", "身份验证信息仅指用户自愿保存的可选 AI 服务 API Key；网络记录仅指用户主动捕获页面的 URL、标题和时间；输入值会被遮罩。")
    add_field(doc, "数据使用", "身份验证信息：仅处理用户自愿配置的可选 AI 服务 API Key，密钥保存在 Chrome 本地存储中，并仅发送给用户选择的 AI 服务商完成其主动请求。网络记录：仅保存用户主动捕获页面的 URL、标题和时间。网站内容：采集可见文本片段、设计令牌、布局指标、资源地址、截图和脱敏证据。用户活动：仅在用户主动发起的采集会话中观察滚动、悬停、聚焦、打开和时间证据，输入值会被遮罩。数据仅用于采集、分析、本地存储、导出和用户主动请求的 AI 生成，不出售、不用于广告、不用于信用评估，也不用于无关用途。", "与 docs/privacy.md 保持一致。")
    add_field(doc, "政策确认", "勾选全部三项：不向第三方出售或传输用户数据；不将用户数据用于与单一用途无关的目的；不将用户数据用于信用评估或贷款目的。", "三项都是发布前必选。")
    add_field(doc, "审核测试说明", "1. 安装标准版扩展，不需要账号。\n2. 打开普通公开 HTTPS 网页，点击工具栏中的 Yotsuba 网页风格提取器，侧边栏应默认打开。\n3. 保持“设计参照”模式，点击“智能捕获”，等待结果出现在侧边栏。\n4. 打开“覆盖”和“历史”查看提取状态。\n5. 导出风格参考包；此流程不需要 AI 密钥。\n6. 仅在获得授权时选择“经授权重建”，确认授权后再测试重建资料。\n扩展不会自动导航、提交表单或执行合成点击；chrome:// 等受限页面会被拒绝。", "完整粘贴到审核测试说明。")

    add_paragraph(doc, "五、提交前检查", bold=True, color="111812", size=15, after=8)
    for item in [
        "只上传 dist/yotsuba-lens-0.3.1-standard-chrome.zip，不上传 Collector 版本。",
        "分发范围选择公开，地区选择所有地区。",
        "审核通过后关闭自动发布或选择暂存发布。",
        "确认图标、截图、小型宣传图块和顶部宣传图块均已上传。",
        "确认隐私权政策、支持网址和说明文字均可公开访问。",
        "确认交易者身份按实际情况填写；当前个人免费开源项目通常为非交易者。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_docx_font(r)
        r.font.size = Pt(10)
    out = ROOT / "docs" / "chrome-web-store-cn-fill-pack.docx"
    doc.save(out)
    return out


make_icon()
make_small_promo()
normalize_screenshots()
make_top_promo()
make_docx()
print("Generated Chrome Web Store assets and Chinese fill pack")
